#!/usr/bin/env python3
"""
Script to parse and extract content from CV-related .docx files
"""

import os
import json
from docx import Document
from pathlib import Path

def extract_text_from_docx(file_path):
    """Extract text content from a .docx file"""
    try:
        doc = Document(file_path)
        text_content = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        return text_content
    except Exception as e:
        return [f"Error reading file: {str(e)}"]

def analyze_cv_content(text_content, filename):
    """Analyze the structure and content of CV files"""
    analysis = {
        "filename": filename,
        "total_lines": len(text_content),
        "sections": [],
        "key_info": {},
        "languages": []
    }

    # Detect language based on filename and content
    if "anglictina" in filename.lower() or "english" in filename.lower():
        analysis["languages"].append("English")
    elif "sk" in filename.lower() or any("č" in line or "š" in line or "ž" in line for line in text_content):
        analysis["languages"].append("Slovak")
    elif "persönliche" in filename.lower() or any("ä" in line or "ö" in line or "ü" in line for line in text_content):
        analysis["languages"].append("German")

    # Identify common CV sections
    section_keywords = {
        "personal_info": ["name", "email", "phone", "address", "date of birth", "personal data"],
        "education": ["education", "university", "school", "degree", "graduation"],
        "experience": ["experience", "work", "employment", "job", "position"],
        "skills": ["skills", "languages", "technical", "competencies"],
        "motivation": ["motivation", "objective", "goal", "motivačný"]
    }

    for section, keywords in section_keywords.items():
        for line in text_content:
            if any(keyword.lower() in line.lower() for keyword in keywords):
                analysis["sections"].append(section)
                break

    # Extract key information
    for line in text_content:
        if "@" in line and "." in line:
            analysis["key_info"]["email"] = line
        elif any(char.isdigit() for char in line) and len(line) > 8:
            if any(word in line.lower() for word in ["phone", "tel", "mobile"]):
                analysis["key_info"]["phone"] = line

    return analysis

def main():
    """Main function to parse all CV files"""
    cv_files = [
        "Motivačný list.docx",
        "Persönliche Daten.docx",
        "CURRICULUM VITAE sk.docx",
        "Curriculum Vitae anglictina.docx"
    ]

    all_analyses = {}

    for filename in cv_files:
        if os.path.exists(filename):
            print(f"\n{'='*50}")
            print(f"Parsing: {filename}")
            print(f"{'='*50}")

            # Extract text content
            text_content = extract_text_from_docx(filename)

            # Analyze content
            analysis = analyze_cv_content(text_content, filename)
            all_analyses[filename] = {
                "analysis": analysis,
                "content": text_content
            }

            # Print analysis
            print(f"Language(s): {', '.join(analysis['languages'])}")
            print(f"Total lines: {analysis['total_lines']}")
            print(f"Identified sections: {', '.join(analysis['sections'])}")
            if analysis['key_info']:
                print("Key information found:")
                for key, value in analysis['key_info'].items():
                    print(f"  {key}: {value}")

            print(f"\nContent preview (first 10 lines):")
            for i, line in enumerate(text_content[:10]):
                print(f"{i+1:2d}: {line}")

            if len(text_content) > 10:
                print(f"... and {len(text_content) - 10} more lines")

    # Save complete analysis to JSON file
    with open("cv_analysis.json", "w", encoding="utf-8") as f:
        json.dump(all_analyses, f, ensure_ascii=False, indent=2)

    print(f"\n{'='*50}")
    print("Complete analysis saved to cv_analysis.json")
    print(f"{'='*50}")

if __name__ == "__main__":
    main()